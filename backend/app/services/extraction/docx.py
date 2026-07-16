"""Word (.docx) extractor using python-docx."""

from __future__ import annotations

from pathlib import Path

import docx

from .base import ExtractedImage, ExtractedTable, ExtractionResult


def extract_docx(path: Path) -> ExtractionResult:
    document = docx.Document(str(path))
    result = ExtractionResult(page_count=1)

    texts = [p.text for p in document.paragraphs if p.text.strip()]
    result.text = "\n".join(texts)

    for table in document.tables:
        rows = [[cell.text for cell in row.cells] for row in table.rows]
        result.tables.append(ExtractedTable(rows=rows))

    # Embedded images live in the package part relationships.
    try:
        for rel in document.part.rels.values():
            if "image" in rel.reltype:
                blob = rel.target_part.blob
                ext = rel.target_part.partname.ext.lstrip(".") or "png"
                result.images.append(ExtractedImage(data=blob, ext=ext))
    except Exception:
        pass

    return result
